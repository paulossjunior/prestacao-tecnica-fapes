"""Extração de texto de arquivos enviados (PDF, DOCX, TXT/MD)."""
import io
import os

from pypdf import PdfReader
from docx import Document

# Limite defensivo de caracteres por documento enviado ao modelo.
MAX_CHARS = 120_000

SUPORTADOS = (".pdf", ".docx", ".txt", ".md")


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    partes = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(partes)


def _from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes.append(" | ".join(c.text for c in linha.cells))
    return "\n".join(partes)


def _from_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    """Extrai texto conforme a extensão do arquivo. Trunca em MAX_CHARS."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        texto = _from_pdf(data)
    elif ext == ".docx":
        texto = _from_docx(data)
    elif ext in (".txt", ".md"):
        texto = _from_text(data)
    else:
        raise ValueError(f"Formato não suportado: {ext} (use pdf, docx, txt ou md)")

    texto = texto.strip()
    if not texto:
        raise ValueError(f"Nenhum texto extraído de {filename} (PDF pode ser escaneado/imagem)")
    if len(texto) > MAX_CHARS:
        texto = texto[:MAX_CHARS] + "\n\n[...documento truncado...]"
    return texto


def extract_many(files) -> str:
    """Concatena o texto de vários uploads (Flask), cada um com cabeçalho de origem."""
    blocos = []
    for f in files:
        if not f or not f.filename:
            continue
        data = f.read()
        if not data:
            continue
        texto = extract_text(f.filename, data)
        blocos.append(f"--- ARQUIVO: {f.filename} ---\n{texto}")
    if not blocos:
        raise ValueError("Nenhum arquivo válido enviado.")
    return "\n\n".join(blocos)


def extract_pairs(pares) -> str:
    """Concatena o texto de vários (nome, bytes). Usado no processamento assíncrono."""
    blocos = []
    for nome, data in pares:
        if not nome or not data:
            continue
        texto = extract_text(nome, data)
        blocos.append(f"--- ARQUIVO: {nome} ---\n{texto}")
    if not blocos:
        raise ValueError("Nenhum arquivo válido enviado.")
    return "\n\n".join(blocos)


def extract_file(path: str) -> str:
    """Lê e extrai o texto de um arquivo do disco."""
    with open(path, "rb") as f:
        return extract_text(path, f.read())


def extract_folder(dirpath: str) -> str:
    """Concatena o texto de todos os arquivos suportados dentro de uma pasta
    (recursivo), cada um com cabeçalho de origem. Ignora não suportados."""
    if not os.path.isdir(dirpath):
        raise ValueError(f"Pasta não encontrada: {dirpath}")

    caminhos = []
    for raiz, _, arquivos in os.walk(dirpath):
        for nome in sorted(arquivos):
            if nome.startswith(".") or nome.startswith("~$"):
                continue  # ocultos / lock temporário do Office
            if os.path.splitext(nome)[1].lower() in SUPORTADOS:
                caminhos.append(os.path.join(raiz, nome))

    if not caminhos:
        raise ValueError(f"Nenhum documento suportado em {dirpath} (pdf/docx/txt/md).")

    blocos = []
    for caminho in sorted(caminhos):
        rel = os.path.relpath(caminho, dirpath)
        texto = extract_file(caminho)
        blocos.append(f"--- ARQUIVO: {rel} ---\n{texto}")
    return "\n\n".join(blocos)
